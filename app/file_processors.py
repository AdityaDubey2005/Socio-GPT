# app/file_processors.py
import os
import mimetypes
from typing import Optional, Dict, Any, List
from pathlib import Path
import tempfile

# PDF processing
try:
    import PyPDF2
    import fitz  # PyMuPDF
    HAS_PDF_SUPPORT = True
except ImportError:
    HAS_PDF_SUPPORT = False

# Document processing
try:
    import docx
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False

# Excel/CSV processing
import pandas as pd
import csv

# Image processing for OCR
try:
    import pytesseract
    from PIL import Image
    HAS_OCR_SUPPORT = True
except ImportError:
    HAS_OCR_SUPPORT = False

def process_uploaded_file(file_path: str, file_type: str) -> str:
    """
    Process uploaded file based on its type and extract meaningful content
    """
    try:
        # Determine file type if not provided
        if not file_type:
            file_type = mimetypes.guess_type(file_path)[0] or 'unknown'
        
        file_extension = Path(file_path).suffix.lower()
        
        # Route to appropriate processor
        if file_type.startswith('text/') or file_extension in ['.txt', '.md', '.log']:
            return process_text_file(file_path)
        
        elif file_type == 'application/pdf' or file_extension == '.pdf':
            return process_pdf_file(file_path)
        
        elif file_type.startswith('image/') or file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            return process_image_file_for_text(file_path)
        
        elif file_type in ['text/csv', 'application/csv'] or file_extension == '.csv':
            return process_csv_file(file_path)
        
        elif file_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'] or file_extension in ['.xlsx', '.xls']:
            return process_excel_file(file_path)
        
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_extension == '.docx':
            return process_docx_file(file_path)
        
        elif file_type == 'application/json' or file_extension == '.json':
            return process_json_file(file_path)
        
        else:
            return f"Unsupported file type: {file_type}. File extension: {file_extension}"
    
    except Exception as e:
        return f"Error processing file: {str(e)}"

def process_text_file(file_path: str) -> str:
    """
    Process plain text files
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                
                # Limit content size for processing
                if len(content) > 50000:  # ~50KB limit
                    content = content[:50000] + "\n\n[Content truncated for processing...]"
                
                return f"Text file content:\n\n{content}"
            
            except UnicodeDecodeError:
                continue
        
        return "Could not decode text file with any standard encoding"
    
    except Exception as e:
        return f"Error reading text file: {str(e)}"

def process_pdf_file(file_path: str) -> str:
    """
    Extract text from PDF files using multiple methods
    """
    if not HAS_PDF_SUPPORT:
        return "PDF processing not available. Install PyPDF2 and PyMuPDF for PDF support."
    
    try:
        text_content = []
        
        # Method 1: Try PyMuPDF (fitz) - better for complex PDFs
        try:
            doc = fitz.open(file_path)
            for page_num in range(min(10, len(doc))):  # Limit to first 10 pages
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    text_content.append(f"Page {page_num + 1}:\n{text}")
            doc.close()
        except Exception as e:
            print(f"PyMuPDF failed: {e}")
        
        # Method 2: Fallback to PyPDF2 if PyMuPDF fails
        if not text_content:
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page_num in range(min(10, len(pdf_reader.pages))):
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"Page {page_num + 1}:\n{text}")
            except Exception as e:
                return f"Error extracting PDF text: {str(e)}"
        
        if text_content:
            full_text = '\n\n'.join(text_content)
            # Limit content size
            if len(full_text) > 20000:
                full_text = full_text[:20000] + "\n\n[Content truncated for processing...]"
            
            return f"PDF content extracted:\n\n{full_text}"
        else:
            return "PDF appears to be empty or contains only images/graphics. Consider using OCR for image-based PDFs."
    
    except Exception as e:
        return f"Error processing PDF: {str(e)}"

def process_image_file_for_text(file_path: str) -> str:
    """
    Extract text from images using OCR
    """
    if not HAS_OCR_SUPPORT:
        return "OCR not available. Install pytesseract and Tesseract for text extraction from images."
    
    try:
        # Open image
        image = Image.open(file_path)
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Extract text using OCR
        extracted_text = pytesseract.image_to_string(image, lang='eng')
        
        if extracted_text.strip():
            return f"Text extracted from image:\n\n{extracted_text.strip()}"
        else:
            return "No readable text found in the image."
    
    except Exception as e:
        return f"Error extracting text from image: {str(e)}"

def process_csv_file(file_path: str) -> str:
    """
    Process CSV files and provide summary
    """
    try:
        # Try to read with pandas first
        df = pd.read_csv(file_path, nrows=1000)  # Limit to first 1000 rows
        
        summary_parts = []
        summary_parts.append(f"CSV file summary:")
        summary_parts.append(f"Shape: {df.shape[0]} rows × {df.shape[1]} columns")
        summary_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
        
        # Data types
        summary_parts.append(f"\nColumn data types:")
        for col, dtype in df.dtypes.items():
            summary_parts.append(f"  {col}: {dtype}")
        
        # Sample data
        summary_parts.append(f"\nFirst few rows:")
        summary_parts.append(df.head(3).to_string())
        
        # Basic statistics for numeric columns
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            summary_parts.append(f"\nNumeric column statistics:")
            summary_parts.append(df[numeric_cols].describe().to_string())
        
        return '\n'.join(summary_parts)
    
    except Exception as e:
        return f"Error processing CSV: {str(e)}"

def process_excel_file(file_path: str) -> str:
    """
    Process Excel files and provide summary
    """
    try:
        # Read Excel file
        excel_file = pd.ExcelFile(file_path)
        sheet_names = excel_file.sheet_names
        
        summary_parts = []
        summary_parts.append(f"Excel file summary:")
        summary_parts.append(f"Number of sheets: {len(sheet_names)}")
        summary_parts.append(f"Sheet names: {', '.join(sheet_names)}")
        
        # Process first sheet or first few sheets
        for sheet_name in sheet_names[:3]:  # Limit to first 3 sheets
            df = pd.read_excel(file_path, sheet_name=sheet_name, nrows=100)
            
            summary_parts.append(f"\nSheet '{sheet_name}':")
            summary_parts.append(f"  Shape: {df.shape[0]} rows × {df.shape[1]} columns")
            summary_parts.append(f"  Columns: {', '.join(df.columns.tolist())}")
            
            # Sample data
            if not df.empty:
                summary_parts.append(f"  Sample data:")
                summary_parts.append("  " + df.head(2).to_string().replace('\n', '\n  '))
        
        return '\n'.join(summary_parts)
    
    except Exception as e:
        return f"Error processing Excel file: {str(e)}"

def process_docx_file(file_path: str) -> str:
    """
    Process Word documents (.docx)
    """
    if not HAS_DOCX_SUPPORT:
        return "DOCX processing not available. Install python-docx for Word document support."
    
    try:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        table_content = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):  # Only add non-empty rows
                    table_content.append(' | '.join(row_text))
        
        # Combine content
        full_content = []
        
        if text_content:
            full_content.append("Document text:")
            full_content.append('\n'.join(text_content))
        
        if table_content:
            full_content.append("\nTables found:")
            full_content.extend(table_content)
        
        if full_content:
            result = '\n\n'.join(full_content)
            # Limit content size
            if len(result) > 20000:
                result = result[:20000] + "\n\n[Content truncated for processing...]"
            return result
        else:
            return "Word document appears to be empty or contains no readable text."
    
    except Exception as e:
        return f"Error processing Word document: {str(e)}"

def process_json_file(file_path: str) -> str:
    """
    Process JSON files and provide summary
    """
    try:
        import json
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        summary_parts = []
        summary_parts.append("JSON file summary:")
        
        if isinstance(data, dict):
            summary_parts.append(f"Type: Dictionary with {len(data)} keys")
            summary_parts.append(f"Keys: {', '.join(list(data.keys())[:10])}")  # Show first 10 keys
            
            # Show structure
            for key, value in list(data.items())[:5]:  # Show first 5 items
                value_type = type(value).__name__
                if isinstance(value, (list, dict)):
                    summary_parts.append(f"  {key}: {value_type} (length: {len(value)})")
                else:
                    value_str = str(value)[:100]  # Truncate long values
                    summary_parts.append(f"  {key}: {value_str}")
        
        elif isinstance(data, list):
            summary_parts.append(f"Type: Array with {len(data)} items")
            if data and isinstance(data[0], dict):
                summary_parts.append(f"Sample item keys: {', '.join(list(data[0].keys())[:10])}")
        
        else:
            summary_parts.append(f"Type: {type(data).__name__}")
            summary_parts.append(f"Content: {str(data)[:500]}")
        
        return '\n'.join(summary_parts)
    
    except Exception as e:
        return f"Error processing JSON file: {str(e)}"

def get_file_metadata(file_path: str) -> Dict[str, Any]:
    """
    Get metadata about the uploaded file
    """
    try:
        file_stat = os.stat(file_path)
        file_path_obj = Path(file_path)
        
        metadata = {
            'filename': file_path_obj.name,
            'extension': file_path_obj.suffix,
            'size_bytes': file_stat.st_size,
            'size_mb': round(file_stat.st_size / (1024 * 1024), 2),
            'mime_type': mimetypes.guess_type(file_path)[0],
            'is_text_based': is_text_based_file(file_path),
            'processing_supported': is_processing_supported(file_path)
        }
        
        return metadata
    
    except Exception as e:
        return {'error': str(e)}

def is_text_based_file(file_path: str) -> bool:
    """
    Determine if file is primarily text-based
    """
    text_extensions = {'.txt', '.md', '.csv', '.json', '.xml', '.html', '.css', '.js', '.py', '.java', '.cpp', '.c', '.h'}
    text_mime_types = {'text/', 'application/json', 'application/xml'}
    
    extension = Path(file_path).suffix.lower()
    mime_type = mimetypes.guess_type(file_path)[0] or ''
    
    return (extension in text_extensions or 
            any(mime_type.startswith(t) for t in text_mime_types))

def is_processing_supported(file_path: str) -> bool:
    """
    Check if file type is supported for processing
    """
    supported_extensions = {'.txt', '.pdf', '.csv', '.xlsx', '.xls', '.docx', '.json', '.md', '.jpg', '.jpeg', '.png', '.bmp', '.tiff'}
    extension = Path(file_path).suffix.lower()
    return extension in supported_extensions

def cleanup_temp_files(file_paths: List[str]):
    """
    Clean up temporary files
    """
    for file_path in file_paths:
        try:
            if os.path.exists(file_path) and '/temp_uploads/' in file_path:
                os.unlink(file_path)
        except Exception as e:
            print(f"Warning: Could not delete temp file {file_path}: {e}")